"""
File handling utilities for document processing.
"""

import io
from pathlib import Path
from typing import Tuple, Optional
import logging

logger = logging.getLogger(__name__)


def get_file_type(filename: str) -> str:
    """
    Determine the file type from filename extension.
    
    Args:
        filename: Name of the file
    
    Returns:
        File type string ('pdf', 'docx', or 'unknown')
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".docx":
        return "docx"
    elif ext == ".doc":
        return "doc"
    return "unknown"


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, dict]:
    """
    Extract text from PDF file content.
    
    Args:
        file_content: Raw PDF file bytes
    
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    try:
        import PyPDF2
        
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        metadata = {
            "page_count": len(reader.pages),
            "has_images": False
        }
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title", "")
            metadata["author"] = reader.metadata.get("/Author", "")
        
        logger.info(f"Extracted {len(text_parts)} pages from PDF")
        return full_text, metadata
        
    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        raise ImportError("PyPDF2 is required for PDF processing")
    except Exception as e:
        logger.error(f"Error extracting PDF text: {str(e)}")
        raise


def extract_text_from_docx(file_content: bytes) -> Tuple[str, dict]:
    """
    Extract text from DOCX file content.
    
    Args:
        file_content: Raw DOCX file bytes
    
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    try:
        from docx import Document
        
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        text_parts = []
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "has_tables": len(doc.tables) > 0
        }
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        core_props = doc.core_properties
        metadata["title"] = core_props.title or ""
        metadata["author"] = core_props.author or ""
        
        logger.info(f"Extracted text from DOCX with {len(doc.paragraphs)} paragraphs")
        return full_text, metadata
        
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise ImportError("python-docx is required for DOCX processing")
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {str(e)}")
        raise


def extract_text(file_content: bytes, filename: str) -> Tuple[str, dict]:
    """
    Extract text from file based on file type.
    
    Args:
        file_content: Raw file bytes
        filename: Name of the file
    
    Returns:
        Tuple of (extracted_text, metadata_dict)
    
    Raises:
        ValueError: If file type is not supported
    """
    file_type = get_file_type(filename)
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_content)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
